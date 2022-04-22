from tkinter import *
from tkcalendar import *
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from datetime import date,datetime
from cryptography.fernet import Fernet
from docx import Document

import tkinter.messagebox as tm
import win32api
import os
import re
import sqlite3


def main():
    window=Tk()
    window.geometry('1350x730')
    window.title("Maruti Optics")
    window.resizable(False,False)
    window.iconbitmap(False,"logo.ico")
    

    img = PhotoImage(file="bg1.png")
    label = Label(window,image=img,height=730,width=1366)
    label.place(x=0,y=0)

    try:
        if domain == 'true':
            num = 0000
    except:
        num = 2000
        
    window.after(num,window.destroy)
    window.mainloop()
    
    if not os.path.exists('mosecure.txt'):
            admin()
    else:
        login()
    
def login():
    global window1
    window1 = Tk()
    window1.configure(background="lightblue")#WINDOW BACKGROUND

    window1.geometry('1350x730')
    window1.title("Maruti Optics")
    window1.resizable(False,False)
    

    global passw, pass_entry
    passw = StringVar()

    teacher_label=Label(window1,anchor=W,text=' Login',font=("bold",13),height=2,width=150,fg='black')
    teacher_label.place(x=0,y=50)

    des_button= Button(window1,text='Exit',font=("bold",18),bd=0,height=1,fg='gray',command=window1.destroy)                       
    des_button.place(x=1250,y=50)

    pass_label=Label(window1,text="Enter Password",font=("bold",13),fg='blue',bg='lightblue')
    pass_label.place(x=480,y=310)
    pass_entry=Entry(window1,width=25,font=("bold",20),textvariable=passw,bg='pink',show='*')
    pass_entry.place(x=480,y=335)

    login_button= Button(window1,text='LOG IN',font=("bold",15),width=15,bg='black',fg='yellow',command=verify1)                        
    login_button.place(x=590,y=380)

    window1.mainloop()
    

def gen():
    # key generation
    key = Fernet.generate_key()   
    ## string the key in a file
    with open('filekey.key', 'wb') as filekey:
       filekey.write(key)
       
def call():
    return open("filekey.key", "rb").read()


def verify1():

    pin = passw.get()
    pass_entry.delete(0,"end")

    key = call()
    fernet = Fernet(key)
    
    # opening the encrypted file
    with open('mosecure.txt', 'rb') as enc_file:
        encrypted = enc_file.read()
      
    # decrypting the file
    decrypted = fernet.decrypt(encrypted).decode()
 
    if pin == decrypted :
        success = tm.showinfo("Login Success","You Have Successfully Logged In!\n               Welcome.")
        if success:
            window1.destroy()
            admin()
    else:
        tm.showinfo("Login Failed","Your Password is Incorrect!\n        Please Try Again.")
        
def close1():
    global domain
    domain = 'true'
    window2.destroy()
    main()

   
def admin():
    
    global window2
    global justice
    justice = 'true'
    window2=Tk() #TKINTER WINDOW
    window2.configure(background="lightblue")#WINDOW BACKGROUND

    window2.geometry('1350x730')
    window2.title("Maruti Optics")
    window2.resizable(False,False)

    img = PhotoImage(file="bg3.png")
    label = Label(window2,image=img,height=730,width=1366)
    label.place(x=0,y=0)
    
    title_label=Label(window2,anchor=CENTER,text='Maruti Optics',font=("Cosmic Sans MS",13,"bold"),height=1,width=50,fg='black')#,bg='orange')
    title_label.place(x=420,y=12)

    text1='  ADMIN'
    course_label=Label(window2,anchor=W,text=text1,font=("Cosmic Sans MS",13,"bold"),height=2,width=150,fg='black')
    course_label.place(x=0,y=50)

    menubutton = Menubutton(window2, text="Settings",font=("Cosmic Sans MS",15,"bold"))#,bg='orange')
    # Create pull down menu
    mainmenu = Menu(menubutton, tearoff = 0)
    menubutton["menu"] = mainmenu

    submenu = Menu(mainmenu,tearoff = 0)
    mainmenu.add_cascade(label="Security",menu = submenu)
  
    if os.path.exists('mosecure.txt'):
        submenu.add_command(label="Change Password",command=change_password)
        submenu.add_command(label="Remove Password",command=rm_password)
        mainmenu.add_separator()
        mainmenu.add_command(label="Logout",command=close1)

    else:
        # Add some commands
        submenu.add_command(label="Add Password",command=add_password)
        mainmenu.add_separator()
        mainmenu.add_command(label="Exit",command=window2.destroy)

    menubutton.place(x=1240,y=55)

    cust_button= Button(window2,text='Patients',font=("bold",18),width=20,bg='black',fg='yellow',command=Customer)                       
    cust_button.place(x=540,y=280)

    print_button= Button(window2,text='View Patients',font=("bold",18),width=20,bg='black',fg='yellow',command=view0)
    print_button.place(x=540,y=360)

    att_button= Button(window2,text='Receipt',font=("bold",18),width=20,bg='black',fg='yellow',command=print0)                        
    att_button.place(x=540,y=440)

    window2.mainloop()

def view0():
    global justice
    justice = 'true'
    view()
    
def print0():
    global justice
    justice = 'false'
    view()

def add_password():

    global window3
    window3 = Toplevel()
    window3.configure(background="lightblue")#WINDOW BACKGROUND

    window3.geometry('550x350')
    window3.resizable(False,False)
    window3.title("Security")

    def backW():
        pass_label.config(text='Enter Password')
        addpass.config(textvariable=passw)
        addpass.delete(0,"end")

        conf_button.config(text='Next',command=nextW)
        b_button.config(bd=0,bg='lightblue',fg='lightblue',command='')

        
    def nextW():
        pass_label.config(text='Confirm Password')
        addpass.config(textvariable=passw2)
        addpass.delete(0,"end")

        conf_button.config(text='Confirm',command=encrypt_password)
        b_button.config(text='Previous',bd=2,bg='black',fg='yellow',command=backW)
        b_button['state'] = NORMAL


    global passw
    global passw2
    global addpass
    
    passw = StringVar()
    passw2 = StringVar()

    pass_label=Label(window3,text="Enter Password",font=("bold",13),fg='blue',bg='lightblue')
    pass_label.place(x=30,y=90)
    addpass=Entry(window3,width=20,font=("bold",20),textvariable=passw,bg='pink',show='*')
    addpass.place(x=30,y=120)

    conf_button= Button(window3,text='Next',font=("bold",15),width=10,bg='black',fg='yellow',command=nextW)
    conf_button.place(x=400,y=220)

    b_button= Button(window3,text='',font=("bold",15),width=10,bd=0,bg='lightblue',fg='lightblue',command='')
    b_button.place(x=30,y=220)
    b_button['state'] = DISABLED

    window3.grab_set()
    window3.mainloop()
    

def encrypt_password():
    gen()
    key = call()
    fernet = Fernet(key)
    
    if passw.get() == passw2.get():
        original=passw2.get()
        encrypted = fernet.encrypt(original.encode())  # encrypting the file
        
        ## opening the file in write mode and writing the encrypted data
        with open('mosecure.txt', 'wb') as encrypted_file:
            encrypted_file.write(encrypted)
        success = tm.showinfo("Security","Your Password is Set!")
        if success:
            window3.destroy()          
            window2.destroy()
            login()  
    else:
        tm.showinfo("Security","Password Don't Match!\n  Please Try Again.")
        addpass.delete(0,"end")
        

def rm_password():
    
    global window4
    window4 = Toplevel()
    window4.configure(background="lightblue")#WINDOW BACKGROUND

    window4.geometry('550x350')
    window4.resizable(False,False)
    window4.title("Security")

    global passw3
    global addpass
    passw3 = StringVar()

    pass_label=Label(window4,text="Enter Password",font=("bold",13),fg='blue',bg='lightblue')
    pass_label.place(x=30,y=90)
    addpass=Entry(window4,width=20,font=("bold",20),textvariable=passw3,bg='pink',show='*')
    addpass.place(x=30,y=120)

    conf_button= Button(window4,text='Submit',font=("bold",15),width=10,bg='black',fg='yellow',command=verify2)
    conf_button.place(x=400,y=220)

    window4.grab_set()
    window4.mainloop()
    

def verify2():
    
    pin = passw3.get()
    addpass.delete(0,"end")
    key = call()
    fernet = Fernet(key)
    # opening the encrypted file
    with open('mosecure.txt', 'rb') as enc_file:
        encrypted = enc_file.read()    
    # decrypting the file
    decrypted = fernet.decrypt(encrypted).decode()
 
    if pin == decrypted :
        window4.destroy()
        sure = tm.askquestion('Security','Are sure you want to remove password?')
        if sure == 'yes':
           if os.path.exists("mosecure.txt"):
               os.remove("mosecure.txt")
               if os.path.exists("filekey.key"):
                   os.remove("filekey.key")
               window2.destroy()
               admin()
           else:
               tm.showinfo('Security','Password Couldn\'t remove.\n Please Restart the software.')
        
    else:
        tm.showinfo("Login Failed","Your Password is Incorrect!\n        Please Try Again.")
        

def change_password():
    global window5
    window5 = Toplevel()
    window5.configure(background="lightblue")#WINDOW BACKGROUND

    window5.geometry('550x350')
    window5.resizable(False,False)
    window5.title("Security")

    def verify():
        pin = passw4.get()
        addpass.delete(0,"end")
        key = call()
        fernet = Fernet(key)
        # opening the encrypted file
        with open('mosecure.txt', 'rb') as enc_file:
            encrypted = enc_file.read()          
        # decrypting the file
        decrypted = fernet.decrypt(encrypted).decode()
     
        if pin == decrypted :
            window5.destroy()
            sure = tm.askquestion('Security','Are sure you want to change password?')
            if sure == 'yes':
                add_password()    
        else:
            tm.showinfo("Login Failed","Your Password is Incorrect!\n        Please Try Again.")
        


    global passw4
    global addpass
    passw4 = StringVar()

    pass_label=Label(window5,text="Enter Current Password",font=("bold",13),fg='blue',bg='lightblue')
    pass_label.place(x=30,y=90)
    addpass=Entry(window5,width=20,font=("bold",20),textvariable=passw4,bg='pink',show='*')
    addpass.place(x=30,y=120)

    conf_button= Button(window5,text='Submit',font=("bold",15),width=10,bg='black',fg='yellow',command=verify)
    conf_button.place(x=400,y=220)

    window5.grab_set()
    window5.mainloop()
    
def Customer():
    def on_entry_click(event):

        t=tot.get()
        d=dis.get()
        a=adv.get()
        
        if t == '' :
            t = 0
        if d == '' :
            d = 0
        if a == '':
            a = 0

        t=float(t)
        d=float(d)
        a=float(a)
        f=t-d-a
        fin_entry.delete(0,"end")
        fin_entry.insert(0, f)
        
    def callback(input1):
           
            if input1 == '' :
                    return True
            elif input1.isdigit() and len(input1) <= 10:                
                    return True
           
            else:        
                    return False       
          
    def callback3(input3):
            try:
                if input3 == '':
                         return True
                if input3 == '-':
                    return True
                f = float(input3)
                return True

            except:
                  return False

    global window6
    window6=Toplevel()#TKINTER WINDOW
    window6.configure(background="lightblue")

    window6.geometry('1350x730')
    window6.resizable(False,False)
    window6.title("CUSTOMER")

    global fname, mno, age

    global sphr, cylr, axr, adr
    global sphl, cyll, axl, adl

    global tot, dis, adv, fin

    global name_entry, mno_entry
    global add_entry, age_entry

    global sr_entry,cr_entry,ar_entry,adr_entry
    global sl_entry,cl_entry,al_entry,adl_entry
    
    global tot_entry,dis_entry,adv_entry,fin_entry    
    
    fname = StringVar()
    mno = StringVar()
    age = StringVar()

    sphr = StringVar()
    cylr = StringVar()
    axr = StringVar()
    adr = StringVar()

    sphl = StringVar()
    cyll = StringVar()
    axl = StringVar()
    adl = StringVar()

    tot = StringVar()
    dis = StringVar()
    adv = StringVar()
    fin = StringVar()

    title_label=Label(window6,anchor=CENTER,text='Maruti Optics',font=("Cosmic Sans MS",13,"bold"),height=1,width=50,fg='black')
    title_label.place(x=420,y=12)

    text2 = ' Customers'
    teacher_label=Label(window6,anchor=W,text=text2,font=("Cosmic Sans MS",13,"bold"),height=2,width=150,fg='black')
    teacher_label.place(x=0,y=50)

    view3_button= Button(window6,text='',font=("bold",14),bd=0,height=1,fg='gray',command='')                       
    view3_button.place(x=155,y=53)

    back_button= Button(window6,text='Exit',font=("bold",18),bd=0,height=1,fg='gray',command=window6.destroy)                       
    back_button.place(x=1250,y=50)

    canvas2 = Frame(window6,height=540,width=1318)
    canvas2.place(x=10,y=110)

    global cal
    cal_label = Label(window6,text="Date:",font=("Cosmic Sans MS",10,"bold"),fg='gray')
    cal_label.place(x=1000,y=130)
    cal = DateEntry(window6, heighth=15,width=18,date_pattern="dd/mm/yyyy")
    cal.place(x=1000,y=150)

    name_label=Label(window6,text="Name:",font=("bold",13),fg='gray')
    name_label.place(x=60,y=220)
    name_entry=Entry(window6,width=28,font=("bold",15),textvariable=fname,bg='pink')
    name_entry.place(x=60,y=245)

    mno_label=Label(window6,text="Mobile No.:",font=("bold",13),fg='gray')
    mno_label.place(x=550,y=220)
    mno_entry=Entry(window6,width=20,font=("bold",15),textvariable=mno,bg='pink')
    mno_entry.place(x=550,y=245)
    reg = window6.register(callback)
    mno_entry.config(validate="key", validatecommand =(reg, '%P'))

    age_label=Label(window6,text="Age:",font=("bold",13),fg='gray')
    age_label.place(x=950,y=220)
    age_entry=Entry(window6,width=20,font=("bold",15),textvariable=age,bg='pink')
    age_entry.place(x=950,y=245)
    age_entry.config(validate="key", validatecommand =(reg, '%P'))

    add_label=Label(window6,text=" Address:",font=("bold",13),fg='gray')
    add_label.place(x=58,y=300)
    add_entry=Text(window6,height=4,width=30,font=("bold",15),bg='pink')
    add_entry.place(x=60,y=325)

    frame2 = Frame(canvas2)
    frame2.place(x=50,y= 340)

    reye_label=Label(frame2,text=" R(Eye) ",font=("bold",13),fg='gray')
    reye_label.grid(row=1,column=2)
    leye_label=Label(frame2,text=" L(Eye) ",font=("bold",13),fg='gray')
    leye_label.grid(row=1,column=3)
    sph_label=Label(frame2,text=" SPH  ",font=("bold",13),fg='gray')
    sph_label.grid(row=2,column=1,sticky=W)
    cyl_label=Label(frame2,text=" CYL  ",font=("bold",13),fg='gray')
    cyl_label.grid(row=3,column=1,sticky=W)
    ax_label=Label(frame2,text=" Axis  ",font=("bold",13),fg='gray')
    ax_label.grid(row=4,column=1,sticky=W)
    add_label=Label(frame2,text=" Add.  ",font=("bold",13),fg='gray')
    add_label.grid(row=5,column=1,sticky=W)

    sr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=sphr)
    sr_entry.grid(row=2,column=2)
    cr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=cylr)
    cr_entry.grid(row=3,column=2)
    ar_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=axr)
    ar_entry.grid(row=4,column=2)
    adr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=adr)
    adr_entry.grid(row=5,column=2)

    sl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=sphl)
    sl_entry.grid(row=2,column=3)
    cl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=cyll)
    cl_entry.grid(row=3,column=3)
    al_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=axl)
    al_entry.grid(row=4,column=3)
    adl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=adl)
    adl_entry.grid(row=5,column=3)

    reg3 = window6.register(callback3)
    sr_entry.config(validate="key", validatecommand =(reg3, '%P'))
    cr_entry.config(validate="key", validatecommand =(reg3, '%P'))
    ar_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adr_entry.config(validate="key", validatecommand =(reg3, '%P'))

    sl_entry.config(validate="key", validatecommand =(reg3, '%P'))
    cl_entry.config(validate="key", validatecommand =(reg3, '%P'))
    al_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adl_entry.config(validate="key", validatecommand =(reg3, '%P'))

    global cal2
    cal2_label = Label(window6,text="Delivery Date:",font=("Cosmic Sans MS",10,"bold"),fg='gray')
    cal2_label.place(x=1000,y=350)
    cal2 = DateEntry(window6, heighth=15,width=18,date_pattern="dd/mm/yyyy")
    cal2.place(x=1000,y=375)

    frame3 = Frame(canvas2)
    frame3.place(x=875,y= 360)

    tot_label=Label(frame3,text=" Total: ",font=("bold",13),fg='gray')
    tot_label.grid(row=1,column=1,sticky=W)
    dis_label=Label(frame3,text=" Discount: ",font=("bold",13),fg='gray')
    dis_label.grid(row=2,column=1,sticky=W)
    adv_label=Label(frame3,text=" Advance: ",font=("bold",13),fg='gray')
    adv_label.grid(row=3,column=1,sticky=W)
    fin_label=Label(frame3,text=" Final Amount: ",font=("bold",13),fg='gray')
    fin_label.grid(row=4,column=1,sticky=W)

    tot_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=tot)
    tot_entry.grid(row=1,column=2)
    dis_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=dis)
    dis_entry.grid(row=2,column=2)
    adv_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=adv)
    adv_entry.grid(row=3,column=2)
    fin_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=fin)
    fin_entry.grid(row=4,column=2)

    tot_entry.config(validate="key", validatecommand =(reg3, '%P'))
    dis_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adv_entry.config(validate="key", validatecommand =(reg3, '%P'))
    fin_entry.config(validate="key", validatecommand =(reg3, '%P'))

    tot_entry.bind('<FocusIn>', on_entry_click)
    dis_entry.bind('<FocusIn>', on_entry_click)
    adv_entry.bind('<FocusIn>', on_entry_click)

    button = Button(window6,text="Submit",font=("bold",15),width=15,bg='black',fg='yellow',command=addpatient)
    button.place(x=570,y=670)

    window6.grab_set() 
    window6.mainloop()
    
def addpatient():
    def addpatient_save():
        base=sqlite3.connect('MarutiOptics.db')
        with base:
            c=base.cursor()      
        query = 'CREATE TABLE IF NOT EXISTS Patient (ID INTEGER PRIMARY KEY,Name TEXT,Mobile_NO TEXT,Age TEXT,Address TEXT,Sph_r TEXT,Cyl_r TEXT,Axis_r TEXT,Add_r TEXT,\
Sph_l TEXT,Cyl_l TEXT,Axis_l TEXT,Add_l TEXT,Total TEXT,Discount TEXT, Advance TEXT,Final_Amount TEXT,Date TEXT,Delivery TEXT)'
        c.execute(query)
        
        c.execute("INSERT INTO Patient (Name,Mobile_NO,Age,Address,Sph_r,Cyl_r,Axis_r,Add_r,Sph_l,Cyl_l,Axis_l,Add_l,Total,Discount,Advance,Final_Amount,Date,Delivery)\
VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(name,mob,vay,add,sphr2,cylr2,axr2,adr2,sphl2,cyll2,axl2,adl2,tot2,dis2,adv2,fin2,ondate,ondate2,))
        base.commit()
        name_entry.delete(0,"end")
        mno_entry.delete(0,"end")
        age_entry.delete(0,"end")
        add_entry.delete('1.0',"end")

        sr_entry.delete(0,"end")
        cr_entry.delete(0,"end")
        ar_entry.delete(0,"end")
        adr_entry.delete(0,"end")

        sl_entry.delete(0,"end")
        cl_entry.delete(0,"end")
        al_entry.delete(0,"end")
        adl_entry.delete(0,"end")

        tot_entry.delete(0,"end")
        dis_entry.delete(0,"end")
        adv_entry.delete(0,"end")
        fin_entry.delete(0,"end")

        tm.showinfo('Customer','Data has been saved.')

    ondate = cal.get_date().strftime("%d/%m/%Y")
    ondate2 = cal2.get_date().strftime("%d/%m/%Y")

    name = fname.get()
    mob = mno.get()
    vay = age.get()
    add = add_entry.get("1.0", "end-1c")

    sphr2 = sphr.get()
    cylr2 = cylr.get()
    axr2 = axr.get()
    adr2 = adr.get()

    sphl2 = sphl.get()
    cyll2 = cyll.get()
    axl2 = axl.get()
    adl2 = adl.get()

    tot2 = tot.get()
    dis2 = dis.get()
    adv2 = adv.get()
    fin2 = fin.get()


    if name == '':
        tm.showinfo('Name','Please Enter Name.')
    elif mob == '':
        tm.showinfo('Mobile','Please Enter Mobile Number.')
    elif fin2 == '':
        tm.showinfo('Final','Please Enter Final Amount')
    elif ondate2 < ondate:
        tm.showinfo('Date','Select delivery date ahead of current date.')
    else:
        addpatient_save()


def print_receipt(id):
    base=sqlite3.connect('MarutiOptics.db')
    with base:
        c=base.cursor()

    row=c.execute("SELECT * FROM patient WHERE id='{}'".format(id))
    s = row.fetchone()

    document = Document("MarutiOptics.docx")

    p = document.add_paragraph('                                                                                                      \
                                      Date: '+s[17])
    p.add_run('\n\nName: ').bold = True
    p.add_run(s[1])
    p.add_run('\nMobile No.: ').bold = True
    p.add_run(s[2])
    p.add_run('\nAge: ').bold = True
    p.add_run(s[3])
    p.add_run('\nAddress: ').bold = True
    p.add_run(s[4])
    p.add_run('\n')

    records = (
        ('SPH ', s[5] , s[9]),
        ('CYL ', s[6] , s[10]),
        ('Axis ', s[7] , s[11]),
        ('Add ', s[8] , s[12]) )    

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '            '
    hdr_cells[1].text = 'R(Eye)'
    hdr_cells[2].text = '       L(Eye)'
    for eye, re, le in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(eye)+"         "
        row_cells[1].text = re
        row_cells[2].text = "           "+le
        
    p = document.add_paragraph('')
    p.add_run('\n\nTotal: ').bold = True
    p.add_run(s[13])
    p.add_run('\nDiscount: ').bold = True
    p.add_run(s[14])
    p.add_run('\nAdvance: ').bold = True
    p.add_run(s[15])
    p.add_run('\nFinal Amount: ').bold = True
    p.add_run(s[16])
    p.add_run('\n')

    p=document.add_paragraph('')
    p.add_run('\nDelivery Date: ').bold = True
    p.add_run(s[18])

    p=document.add_paragraph()
    p.add_run('\n\n\n\n---------------------------------------------------  Thank You ------------------------------------------------').bold = True
    
    try:
        file = 'Patient_'+str(s[0])+'_'+s[1]+'.docx'
        document.save(file)
        
        global window9
        window9 = Toplevel()
        
        window9.title('Print Receipt')
        window9.geometry("650x730")
        window9.resizable(False,False)

        textEditor = Text(window9, width=68, height=33,font=('bold',13))
        textEditor.place(x=15,y=20)
        textEditor.delete(1.0,END)

        document = Document(file)        

        textEditor.insert(END,'\nMaruti Optics\nContact Lens Clinic')
        textEditor.insert(END,'\nShop no.2 Vastu Darshan Building., Below Central Bank, \nNear BMC Ward Office, Azad Nagar, Andheri (E), Mumbai 400069\
\nPhone No: 9969110029/9664289489')
        textEditor.insert(END,'\n____________________________________________________________________')
        textEditor.insert(END,'\n\n                                                                                              Date: '+s[17])
        textEditor.insert(END,'\nName: '+ s[1])
        textEditor.insert(END,'\nMobile No.: '+s[2])
        textEditor.insert(END,'\nAge: '+s[3])
        textEditor.insert(END,'\nAddress: '+s[4])
        textEditor.insert(END,"\n\n")

        for table in document.tables:
            for row in table.rows: 
                for cell in row.cells: 
                    for paragraph in cell.paragraphs:
                        dt=paragraph.text#+"       "
                        textEditor.insert(END,dt)
                textEditor.insert(END,"\n ")
            
        textEditor.insert(END,'\n\nTotal: '+ s[13])
        textEditor.insert(END,'\nDiscount: '+s[14])
        textEditor.insert(END,'\nAdvance: '+s[15])
        textEditor.insert(END,'\nFinal Amount: '+s[16])
        textEditor.insert(END,"\n\nDelivery Date:"+s[18])
        textEditor.insert(END,'\n\n\n\n--------------------------------------------- Thank You ------------------------------------------')
        textEditor.config(state=DISABLED)
        
        def print_file():
            win32api.ShellExecute(0, "print", file, None, ".", 0)

        Button(window9, text="Print",font=("bold",15),width=12,bg='black',fg='yellow',command=print_file).place(x=260,y=656)

        window9.grab_set()
        window9.mainloop()
      
    except:
        tm.showinfo('Error','Please Close the Word Document!')


def delete_data(id):
    
    query=("Delete FROM patient WHERE id={}".format(id))
    base=sqlite3.connect('MarutiOptics.db')
    with base:
        c=base.cursor()

    row=c.execute(query)
    base.commit()
    window7.destroy()
    view()

   
def edit_data(id): # display to edit and update record
    
    def my_update(): # update record
        e5_add=e5.get("1.0", "end-1c")
        ondate = cal.get_date().strftime("%d/%m/%Y")
        ondate2 = cal2.get_date().strftime("%d/%m/%Y")
        data=(e2_name.get(),e3_mob.get(),e4_age.get(),e5_add,e6_sphr.get(),e7_cylr.get(),e8_axisr.get(),e9_adr.get(),e10_sphl.get(),e11_cyll.get(),e12_axisl.get(),
              e13_adl.get(),e15_tot.get(),e16_dis.get(),e17_adv.get(),e18_fin.get(),ondate,ondate2,e1_id.get()) 

        id=c.execute("UPDATE patient SET name='{0[0]}',mobile_no='{0[1]}',age='{0[2]}',Address='{0[3]}',sph_r='{0[4]}',cyl_r='{0[5]}',axis_r='{0[6]}',add_r='{0[7]}',\
Sph_l='{0[8]}',Cyl_l='{0[9]}',Axis_l='{0[10]}',Add_l='{0[11]}',Total='{0[12]}',Discount='{0[13]}',Advance='{0[14]}',Final_Amount='{0[15]}',Date='{0[16]}',\
Delivery='{0[17]}' WHERE id={0[18]}".format(data)) 
        base.commit()

        window8.destroy()
        display()

        
    def callback(input1):     
            if input1 == '' :
                    return True
            elif input1.isdigit() and len(input1) <= 10:                
                    return True    
            else:        
                    return False       
          
    def callback3(input3):
            try:
                if input3 == '':
                         return True
                if input3 == '-':
                    return True
                f = float(input3)
                return True
            except:
                  return False
    
    global window8
    window8=Toplevel()#TKINTER WINDOW
    window8.configure(background="lightblue")

    window8.geometry('800x530')
    window8.resizable(False,False)
    window8.title("UPDATE CUSTOMER")
    
    global i # start row after the last line of display
    #collect record based on id and present for updation.
    base=sqlite3.connect('MarutiOptics.db')
    with base:
        c=base.cursor()

    row=c.execute("SELECT * FROM patient WHERE id='{}'".format(id))
    s = row.fetchone() # row details as tuple

    row2=c.execute("SELECT (SELECT COUNT(*) FROM Patient AS t2 WHERE t2.id <= t1.id) AS row_Num FROM Patient AS t1 where id ='{}' ORDER BY id;".format(id))
    s2=row2.fetchone()      

    e1_id=StringVar() # String variable 
    e2_name=StringVar()
    e3_mob=StringVar()
    e4_age=StringVar()
    e5_add=StringVar()

    e6_sphr=StringVar()
    e7_cylr=StringVar()
    e8_axisr=StringVar()
    e9_adr=StringVar()

    e10_sphl=StringVar()
    e11_cyll=StringVar()
    e12_axisl=StringVar()
    e13_adl=StringVar()

    e15_tot=StringVar()
    e16_dis=StringVar()
    e17_adv=StringVar()
    e18_fin=StringVar()

    e1_id.set(s[0]) # Set entries 
    e2_name.set(s[1])
    e3_mob.set(s[2])
    e4_age.set(s[3])
    
    e6_sphr.set(s[5])
    e7_cylr.set(s[6])
    e8_axisr.set(s[7])
    e9_adr.set(s[8])

    e10_sphl.set(s[9])
    e11_cyll.set(s[10])
    e12_axisl.set(s[11])
    e13_adl.set(s[12])
    
    e15_tot.set(s[13])
    e16_dis.set(s[14])
    e17_adv.set(s[15])
    e18_fin.set(s[16]) 

#Entries
    frame0 = Canvas(window8,bg='lightblue',highlightbackground='black',highlightthickness='1',height=200,width=350)
    frame0.place(x=15,y=80)
    
    frame = Frame(frame0,bg='lightblue')#,highlightbackground='black',highlightthickness='1')
    frame.place(x=15,y=10)
    text="    Sr.no.: " +str(s2[0])
    e1_label=Label(window8,anchor=W,text=text,font=("bold",14),height=1,width=150,fg='black')
    e1_label.place(x=0,y=30)
    
    back_button= Button(window8,text='Exit',font=("Cosmic Sans MS",11,"bold"),bd=0,fg='gray',command=window8.destroy)                       
    back_button.place(x=740,y=30)
    
    e2_label=Label(frame,text="Name: ",font=("bold",14),fg='Black',bg='lightblue')
    e2_label.grid(row=0,column=0,sticky=W)
    e2=Entry(frame,textvariable=e2_name,width=14,font=("bold",14))
    e2.grid(row=0,column=1)#place(x=15,y=125)#pack()#grid(row=i,column=1)

    e3_label=Label(frame,text="Mobile No: ",font=("bold",14),fg='Black',bg='lightblue')
    e3_label.grid(row=1,column=0,sticky=W)
    e3=Entry(frame,textvariable=e3_mob,width=14,font=("bold",14))
    e3.grid(row=1,column=1)#place(x=15,y=145)#pack()#grid(row=i,column=2)
    reg = window8.register(callback)
    e3.config(validate="key", validatecommand =(reg, '%P'))

    e4_label=Label(frame,text="Age: ",font=("bold",14),fg='Black',bg='lightblue')
    e4_label.grid(row=2,column=0,sticky=W)
    e4=Entry(frame,textvariable=e4_age,width=14,font=("bold",14))
    e4.grid(row=2,column=1)#place(x=240,y=640)#pack()#grid(row=i,column=3)
    e4.config(validate="key", validatecommand =(reg, '%P'))

    e5_label=Label(frame,text="Address: ",font=("bold",14),fg='Black',bg='lightblue')
    e5_label.grid(row=3,column=0,sticky=NW)
    e5=Text(frame,width=14,font=("bold",14),height=4)
    e5.insert(INSERT,s[4])
    e5.grid(row=3,column=1)#place(x=315,y=640)#pack()#grid(row=i,column=4)

#EyeData
    frame2 = Frame(window8,bg='lightblue')
    frame2.place(x=15,y=300)

    reye_label=Label(frame2,text=" R(Eye) ",font=("bold",13),fg='Black',bg='lightblue')
    reye_label.grid(row=1,column=2)
    leye_label=Label(frame2,text=" L(Eye) ",font=("bold",13),fg='Black',bg='lightblue')
    leye_label.grid(row=1,column=3)
    sph_label=Label(frame2,text=" SPH  ",font=("bold",13),fg='Black',bg='lightblue')
    sph_label.grid(row=2,column=1,sticky=W)
    cyl_label=Label(frame2,text=" CYL  ",font=("bold",13),fg='Black',bg='lightblue')
    cyl_label.grid(row=3,column=1,sticky=W)
    ax_label=Label(frame2,text=" Axis  ",font=("bold",13),fg='Black',bg='lightblue')
    ax_label.grid(row=4,column=1,sticky=W)
    add_label=Label(frame2,text=" Add.  ",font=("bold",13),fg='Black',bg='lightblue')
    add_label.grid(row=5,column=1,sticky=W)

    sr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e6_sphr)
    sr_entry.grid(row=2,column=2)
    cr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e7_cylr)
    cr_entry.grid(row=3,column=2)
    ar_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e8_axisr)
    ar_entry.grid(row=4,column=2)
    adr_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e9_adr)
    adr_entry.grid(row=5,column=2)

    sl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e10_sphl)
    sl_entry.grid(row=2,column=3)
    cl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e11_cyll)
    cl_entry.grid(row=3,column=3)
    al_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e12_axisl)
    al_entry.grid(row=4,column=3)
    adl_entry=Entry(frame2,width=10,font=("bold",15),bg='pink',textvariable=e13_adl)
    adl_entry.grid(row=5,column=3)

    reg3 = window8.register(callback3)
    sr_entry.config(validate="key", validatecommand =(reg3, '%P'))
    cr_entry.config(validate="key", validatecommand =(reg3, '%P'))
    ar_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adr_entry.config(validate="key", validatecommand =(reg3, '%P'))

    sl_entry.config(validate="key", validatecommand =(reg3, '%P'))
    cl_entry.config(validate="key", validatecommand =(reg3, '%P'))
    al_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adl_entry.config(validate="key", validatecommand =(reg3, '%P'))

#Calendar
    global cal

    cal_label = Label(window8,text="Date:",font=("Cosmic Sans MS",10,"bold"),fg='Black',bg='lightblue')
    cal_label.place(x=500,y=100)
    cal = DateEntry(window8, heighth=15,width=18,date_pattern="dd/mm/yyyy")

    dt=(datetime.strptime(s[17],"%d/%m/%Y"))#.strftime("%d/%m/%Y")
    cal.set_date(dt)
    cal.place(x=500,y=120)
    global cal2
    cal2_label = Label(window8,text="Delivery Date:",font=("Cosmic Sans MS",10,"bold"),fg='Black',bg='lightblue')
    cal2_label.place(x=500,y=160)
    cal2 = DateEntry(window8, heighth=15,width=18,date_pattern="dd/mm/yyyy")

    dt2=(datetime.strptime(s[18],"%d/%m/%Y"))#.strftime("%d/%m/%Y")
    cal2.set_date(dt2)
    cal2.place(x=500,y=180)


#Payment
    frame3 = Frame(window8,bg='lightblue')
    frame3.place(x=500,y=320)

    tot_label=Label(frame3,text=" Total: ",font=("bold",13),fg='Black',bg='lightblue')
    tot_label.grid(row=1,column=1,sticky=W)
    dis_label=Label(frame3,text=" Discount: ",font=("bold",13),fg='Black',bg='lightblue')
    dis_label.grid(row=2,column=1,sticky=W)
    adv_label=Label(frame3,text=" Advance: ",font=("bold",13),fg='Black',bg='lightblue')
    adv_label.grid(row=3,column=1,sticky=W)
    fin_label=Label(frame3,text=" Final Amount: ",font=("bold",13),fg='Black',bg='lightblue')
    fin_label.grid(row=4,column=1,sticky=W)

    tot_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=e15_tot)
    tot_entry.grid(row=1,column=2)
    dis_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=e16_dis)
    dis_entry.grid(row=2,column=2)
    adv_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=e17_adv)
    adv_entry.grid(row=3,column=2)
    fin_entry=Entry(frame3,width=10,font=("bold",15),bg='pink',textvariable=e18_fin)
    fin_entry.grid(row=4,column=2)

    tot_entry.config(validate="key", validatecommand =(reg3, '%P'))
    dis_entry.config(validate="key", validatecommand =(reg3, '%P'))
    adv_entry.config(validate="key", validatecommand =(reg3, '%P'))
    fin_entry.config(validate="key", validatecommand =(reg3, '%P'))

    b2 =Button(window8,text='Update',command=lambda: my_update(),relief='ridge',font=("bold",15),width=15,bg='black',fg='yellow')  
    b2.place(x=300,y=460)

    window8.grab_set()
    window8.mainloop()
    

def view():
    global window7
    window7 = Toplevel()
    window7.geometry('1350x730')
    window7.title("Maruti Optics")
    window7.resizable(False,False)
    
    global my_w
    my_w= Frame(window7)
    my_w.pack(fill=BOTH, expand=1)

    img = PhotoImage(file="bg4.png")
    label = Label(my_w,image=img,height=730,width=1366)
    label.place(x=0,y=0)

    text2 = 'Customers'
    teacher_label=Label(my_w,anchor=W,text=text2,font=("Cosmic Sans MS",13,"bold"),height=2,width=150,fg='black')
    teacher_label.place(x=0,y=50)

    back_button= Button(my_w,text='Exit',font=("bold",18),bd=0,height=1,fg='gray',command=window7.destroy)                       
    back_button.place(x=1250,y=50)

       
    global i    
    base=sqlite3.connect('MarutiOptics.db')
    with base:
        c=base.cursor()

    #create canvas
    canvas = Canvas(my_w,height=590,width=1310)
    canvas.place(x=10,y=110)#pack(fill=BOTH, expand=1)

    #add scrollbar to the canvas
    myscrollbar= Scrollbar(my_w, orient=VERTICAL,command=canvas.yview)
    myscrollbar.pack(side=RIGHT, fill=Y)

    #configure canvas
    canvas.configure(yscrollcommand=myscrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    #create second frame
    frame2 = Frame(canvas)
    
    #add new frame to a window in canvas
    canvas.create_window((0,0), window=frame2, anchor="nw")
    
    global display
    def display():
        try:
            co=c.execute('select Count(*) from patient')
            for x in co:
                if x[0] == 0:
                    text = 'No Customers Added'
                    noC = Label(window7,text=text,font=("Cosmic Sans MS",20,"bold"))
                    noC.place(x=520,y=370)
            my_cursor=c.execute("SELECT id,name,mobile_no,sph_r,cyl_r,Axis_r,sph_l,cyl_l,Axis_l,Date,Delivery,Total FROM patient")
                  
            global i
            i=2

            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Sr.no")  
            l1.grid(row=0,rowspan=2, column=0)
            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Name")  
            l1.grid(row=0,rowspan=2, column=1)
            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Mobile No.")  
            l1.grid(row=0,rowspan=2, column=2)
            
            l1 = Label(frame2,width=38,height=2,relief='ridge', text="Right Eye")  
            l1.grid(row=0, column=3,columnspan=3)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="SPH")  
            l1.grid(row=1, column=3)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="CYL")  
            l1.grid(row=1, column=4)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="AXIS")  
            l1.grid(row=1, column=5)
            
            l1 = Label(frame2,width=38,height=2,relief='ridge', text="Left Eye")  
            l1.grid(row=0, column=6,columnspan=3)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="SPH")  
            l1.grid(row=1, column=6)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="CYL")  
            l1.grid(row=1, column=7)
            l1 = Label(frame2,width=12,height=2,relief='ridge', text="AXIS")  
            l1.grid(row=1, column=8)

            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Date")  
            l1.grid(row=0,rowspan=2, column=9)
            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Delivery \nDate")  
            l1.grid(row=0,rowspan=2, column=10)
            l1 = Label(frame2,width=12,height=5,relief='ridge', text="Total")  
            l1.grid(row=0,rowspan=2, column=11)
         
            for student in my_cursor: 
                for j in range(12):                                                                                                                                                                                                
                    e = Label(frame2,width=12,height=2,relief='ridge', text=student[j], anchor="w")  
                    e.grid(row=i, column=j)
                    e = Label(frame2,width=12,height=2,relief='ridge', text=i-1, anchor="w") 
                    e.grid(row=i, column=0)
                    
                if justice == 'false':
                    e3 = Button(frame2,width=9,text='Receipt',relief='ridge',
                         command=lambda k=student[0]:print_receipt(k))  
                    e3.grid(row=i, column=15)
                else:
                    e = Button(frame2,width=10,text='Edit',relief='ridge',
                         command=lambda k=student[0]:edit_data(k))  
                    e.grid(row=i, column=13)
                    
                    e2 = Button(frame2,width=9,text='Delete',relief='ridge',
                         command=lambda k=student[0]:delete_data(k))  
                    e2.grid(row=i, column=14)
                    
                    e3 = Button(frame2,width=9,text='Receipt',relief='ridge',
                         command=lambda k=student[0]:print_receipt(k))  
                    e3.grid(row=i, column=15)
                i=i+1

        except:
            text = 'No Customers Added'
            noC = Label(window7,text=text,font=("Cosmic Sans MS",20,"bold"))
            noC.place(x=520,y=370)           
        
    display()
    window7.grab_set()
    window7.mainloop()

        
if __name__ == '__main__':
    main()

